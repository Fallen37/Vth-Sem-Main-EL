"""Content Ingestion Service - handles document upload, processing, and embedding."""

from __future__ import annotations

import os
import re
from datetime import datetime
from typing import TYPE_CHECKING, Any, Optional
from uuid import UUID, uuid4

from pydantic import BaseModel, Field
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from config.settings import get_settings
from src.models.document import (
    ContentFilters,
    ContentMetadata,
    Document,
    DocumentORM,
)
from src.models.enums import DocumentStatus


if TYPE_CHECKING:
    import chromadb
    from sentence_transformers import SentenceTransformer

settings = get_settings()


class ProcessingResult(BaseModel):
    """Result of document processing."""

    document_id: UUID
    chunk_count: int
    embedding_status: str = Field(description="SUCCESS, PARTIAL, or FAILED")
    errors: list[str] = Field(default_factory=list)


class TextExtractor:
    """Extracts text from various file formats."""

    @staticmethod
    def extract_from_txt(file_path: str) -> str:
        """Extract text from a TXT file."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from a PDF file using PyPDF2."""
        try:
            import pypdf

            text_parts = []
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("pypdf is required for PDF processing. Install with: pip install pypdf")


    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            import docx

            doc = docx.Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

    @staticmethod
    def extract_from_image(file_path: str) -> str:
        """Extract text from an image using OCR."""
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except ImportError:
            raise ImportError(
                "pytesseract and Pillow are required for image processing. "
                "Install with: pip install pytesseract Pillow"
            )

    @classmethod
    def extract(cls, file_path: str) -> str:
        """Extract text from a file based on its extension."""
        ext = os.path.splitext(file_path)[1].lower()

        extractors = {
            ".txt": cls.extract_from_txt,
            ".pdf": cls.extract_from_pdf,
            ".docx": cls.extract_from_docx,
            ".png": cls.extract_from_image,
            ".jpg": cls.extract_from_image,
            ".jpeg": cls.extract_from_image,
            ".gif": cls.extract_from_image,
            ".bmp": cls.extract_from_image,
            ".tiff": cls.extract_from_image,
        }

        extractor = extractors.get(ext)
        if not extractor:
            raise ValueError(f"Unsupported file format: {ext}")

        return extractor(file_path)


class TextChunker:
    """Chunks text into smaller pieces for embedding."""

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        min_chunk_size: int = 100,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size

    def chunk(self, text: str) -> list[str]:
        """Split text into overlapping chunks."""
        if not text or not text.strip():
            return []

        # Clean the text
        text = re.sub(r"\s+", " ", text).strip()

        if len(text) <= self.chunk_size:
            return [text] if len(text) >= self.min_chunk_size else []

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                last_period = text.rfind(".", start, end)
                last_question = text.rfind("?", start, end)
                last_exclaim = text.rfind("!", start, end)
                break_point = max(last_period, last_question, last_exclaim)

                if break_point > start + self.min_chunk_size:
                    end = break_point + 1

            chunk = text[start:end].strip()
            if len(chunk) >= self.min_chunk_size:
                chunks.append(chunk)

            # Move start with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break

        return chunks



class EmbeddingService:
    """Generates embeddings using sentence-transformers."""

    _instance: Optional["EmbeddingService"] = None
    _model: Optional[Any] = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self):
        if self._model is None:
            from sentence_transformers import SentenceTransformer
            self._model = SentenceTransformer(settings.embedding_model_name)

    def embed(self, texts: list[str]) -> list[list[float]]:
        """Generate embeddings for a list of texts."""
        if not texts:
            return []
        embeddings = self._model.encode(texts, convert_to_numpy=True)
        return embeddings.tolist()

    def embed_single(self, text: str) -> list[float]:
        """Generate embedding for a single text."""
        embedding = self._model.encode([text], convert_to_numpy=True)
        return embedding[0].tolist()


class ChromaDBClient:
    """Client for ChromaDB vector database operations."""

    _instance: Optional["ChromaDBClient"] = None
    _client: Optional[Any] = None
    _collection: Optional[Any] = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self):
        if self._client is None:
            import chromadb
            from chromadb.config import Settings as ChromaSettings

            # Ensure directory exists
            os.makedirs(settings.chroma_persist_directory, exist_ok=True)

            self._client = chromadb.PersistentClient(
                path=settings.chroma_persist_directory,
                settings=ChromaSettings(anonymized_telemetry=False),
            )
            self._collection = self._client.get_or_create_collection(
                name=settings.chroma_collection_name,
                metadata={"hnsw:space": "cosine"},
            )

    @property
    def collection(self) -> Any:
        """Get the ChromaDB collection."""
        return self._collection

    def add_documents(
        self,
        ids: list[str],
        embeddings: list[list[float]],
        documents: list[str],
        metadatas: list[dict],
    ) -> None:
        """Add documents with embeddings to the collection."""
        self._collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=documents,
            metadatas=metadatas,
        )

    def query(
        self,
        query_embedding: list[float],
        n_results: int = 5,
        where: Optional[dict] = None,
    ) -> dict:
        """Query the collection for similar documents."""
        return self._collection.query(
            query_embeddings=[query_embedding],
            n_results=n_results,
            where=where,
            include=["documents", "metadatas", "distances"],
        )

    def query_by_curriculum(
        self,
        query_embedding: list[float],
        grade: Optional[int] = None,
        syllabus: Optional[str] = None,
        subject: Optional[str] = None,
        chapter: Optional[str] = None,
        topic: Optional[str] = None,
        n_results: int = 5,
    ) -> dict:
        """
        Query the collection with curriculum-based filtering.
        
        Args:
            query_embedding: The query embedding vector
            grade: Filter by grade level (5-10)
            syllabus: Filter by syllabus type (cbse, state)
            subject: Filter by subject name
            chapter: Filter by chapter name
            topic: Filter by topic name
            n_results: Number of results to return
            
        Returns:
            Query results with documents, metadatas, and distances
        """
        # Build where clause for curriculum filtering
        conditions = []
        
        if grade is not None:
            conditions.append({"grade": grade})
        if syllabus is not None:
            conditions.append({"syllabus": syllabus})
        if subject is not None:
            conditions.append({"subject": subject})
        if chapter is not None:
            conditions.append({"chapter": chapter})
        if topic is not None:
            conditions.append({"topic": topic})
        
        where = None
        if conditions:
            if len(conditions) == 1:
                where = conditions[0]
            else:
                where = {"$and": conditions}
        
        return self.query(
            query_embedding=query_embedding,
            n_results=n_results,
            where=where,
        )

    def delete_by_document_id(self, document_id: str) -> None:
        """Delete all chunks for a document."""
        self._collection.delete(where={"document_id": document_id})



class ContentIngestionService:
    """Service for uploading and processing educational documents."""

    def __init__(
        self,
        db: AsyncSession,
        embedding_service: Optional[EmbeddingService] = None,
        chroma_client: Optional[ChromaDBClient] = None,
        text_extractor: Optional[TextExtractor] = None,
        chunker: Optional[TextChunker] = None,
    ):
        self.db = db
        self._embedding_service = embedding_service
        self._chroma_client = chroma_client
        self.text_extractor = text_extractor or TextExtractor()
        self.chunker = chunker or TextChunker()

    @property
    def embedding_service(self) -> EmbeddingService:
        """Lazy initialization of embedding service."""
        if self._embedding_service is None:
            self._embedding_service = EmbeddingService()
        return self._embedding_service

    @property
    def chroma_client(self) -> ChromaDBClient:
        """Lazy initialization of ChromaDB client."""
        if self._chroma_client is None:
            self._chroma_client = ChromaDBClient()
        return self._chroma_client

    async def upload_document(
        self,
        file_path: str,
        filename: str,
        metadata: ContentMetadata,
    ) -> UUID:
        """
        Upload a document with metadata.

        Args:
            file_path: Path to the uploaded file
            filename: Original filename
            metadata: Content metadata (grade, syllabus, subject, chapter, topic, tags, content_type)

        Returns:
            Document ID
        """
        import json
        
        # Validate file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file extension
        ext = os.path.splitext(filename)[1].lower()
        supported_extensions = {".txt", ".pdf", ".docx", ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"}
        if ext not in supported_extensions:
            raise ValueError(f"Unsupported file format: {ext}. Supported: {supported_extensions}")

        # Create document record
        doc_id = str(uuid4())
        document = DocumentORM(
            id=doc_id,
            filename=filename,
            content_type=metadata.content_type.value,
            grade=metadata.grade,
            syllabus=metadata.syllabus.value,
            subject=metadata.subject,
            chapter=metadata.chapter,
            topic=metadata.topic,
            tags=json.dumps(metadata.tags) if metadata.tags else None,
            status=DocumentStatus.PENDING.value,
        )

        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)

        return UUID(doc_id)

    async def process_document(
        self,
        document_id: UUID,
        file_path: str,
    ) -> ProcessingResult:
        """
        Process a document: extract text, chunk, embed, and store in ChromaDB.

        Args:
            document_id: ID of the document to process
            file_path: Path to the file

        Returns:
            ProcessingResult with status and chunk count
        """
        import json
        
        errors: list[str] = []
        doc_id_str = str(document_id)

        # Get document from database
        result = await self.db.execute(
            select(DocumentORM).where(DocumentORM.id == doc_id_str)
        )
        document = result.scalar_one_or_none()

        if not document:
            return ProcessingResult(
                document_id=document_id,
                chunk_count=0,
                embedding_status="FAILED",
                errors=["Document not found in database"],
            )

        # Update status to processing
        document.status = DocumentStatus.PROCESSING.value
        await self.db.commit()

        try:
            # Extract text
            text = self.text_extractor.extract(file_path)

            if not text or not text.strip():
                document.status = DocumentStatus.FAILED.value
                await self.db.commit()
                return ProcessingResult(
                    document_id=document_id,
                    chunk_count=0,
                    embedding_status="FAILED",
                    errors=["No text could be extracted from the document"],
                )

            # Chunk text
            chunks = self.chunker.chunk(text)

            if not chunks:
                document.status = DocumentStatus.FAILED.value
                await self.db.commit()
                return ProcessingResult(
                    document_id=document_id,
                    chunk_count=0,
                    embedding_status="FAILED",
                    errors=["Document too short to create meaningful chunks"],
                )

            # Generate embeddings
            embeddings = self.embedding_service.embed(chunks)

            # Parse tags from JSON if present
            tags = []
            if document.tags:
                try:
                    tags = json.loads(document.tags)
                except json.JSONDecodeError:
                    tags = []

            # Prepare metadata for each chunk with curriculum information
            chunk_ids = [f"{doc_id_str}_chunk_{i}" for i in range(len(chunks))]
            metadatas = [
                {
                    "document_id": doc_id_str,
                    "chunk_index": i,
                    "grade": document.grade,
                    "syllabus": document.syllabus,
                    "subject": document.subject,
                    "chapter": document.chapter,
                    "topic": document.topic or "",
                    "tags": ",".join(tags) if tags else "",
                    "content_type": document.content_type,
                    "filename": document.filename,
                }
                for i in range(len(chunks))
            ]

            # Store in ChromaDB
            self.chroma_client.add_documents(
                ids=chunk_ids,
                embeddings=embeddings,
                documents=chunks,
                metadatas=metadatas,
            )

            # Update document status
            document.chunk_count = len(chunks)
            document.processed_at = datetime.utcnow()
            document.status = DocumentStatus.READY.value
            await self.db.commit()

            return ProcessingResult(
                document_id=document_id,
                chunk_count=len(chunks),
                embedding_status="SUCCESS",
                errors=[],
            )

        except Exception as e:
            errors.append(str(e))
            document.status = DocumentStatus.FAILED.value
            await self.db.commit()

            return ProcessingResult(
                document_id=document_id,
                chunk_count=0,
                embedding_status="FAILED",
                errors=errors,
            )


    async def delete_document(self, document_id: UUID) -> bool:
        """
        Delete a document and its embeddings.

        Args:
            document_id: ID of the document to delete

        Returns:
            True if deleted, False if not found
        """
        doc_id_str = str(document_id)

        # Get document from database
        result = await self.db.execute(
            select(DocumentORM).where(DocumentORM.id == doc_id_str)
        )
        document = result.scalar_one_or_none()

        if not document:
            return False

        # Delete from ChromaDB
        try:
            self.chroma_client.delete_by_document_id(doc_id_str)
        except Exception:
            pass  # Continue even if ChromaDB deletion fails

        # Delete from database
        await self.db.delete(document)
        await self.db.commit()

        return True

    async def list_documents(
        self,
        filters: Optional[ContentFilters] = None,
    ) -> list[Document]:
        """
        List documents with optional filters.

        Args:
            filters: Optional filters for grade, syllabus, subject, chapter, topic, tags, content_type

        Returns:
            List of documents matching filters
        """
        import json
        
        query = select(DocumentORM)

        if filters:
            if filters.grade is not None:
                query = query.where(DocumentORM.grade == filters.grade)
            if filters.syllabus is not None:
                query = query.where(DocumentORM.syllabus == filters.syllabus.value)
            if filters.subject is not None:
                query = query.where(DocumentORM.subject == filters.subject)
            if filters.chapter is not None:
                query = query.where(DocumentORM.chapter == filters.chapter)
            if filters.topic is not None:
                query = query.where(DocumentORM.topic == filters.topic)
            if filters.content_type is not None:
                query = query.where(DocumentORM.content_type == filters.content_type.value)

        result = await self.db.execute(query)
        documents = result.scalars().all()

        # Post-filter by tags if specified (since tags are stored as JSON)
        filtered_docs = []
        for doc in documents:
            if filters and filters.tags:
                doc_tags = []
                if doc.tags:
                    try:
                        doc_tags = json.loads(doc.tags)
                    except json.JSONDecodeError:
                        doc_tags = []
                # Check if any filter tag matches document tags
                if not any(tag in doc_tags for tag in filters.tags):
                    continue
            
            # Parse tags for the document
            tags = []
            if doc.tags:
                try:
                    tags = json.loads(doc.tags)
                except json.JSONDecodeError:
                    tags = []
            
            filtered_docs.append(
                Document(
                    id=UUID(doc.id),
                    filename=doc.filename,
                    content_type=doc.content_type,
                    grade=doc.grade,
                    syllabus=doc.syllabus,
                    subject=doc.subject,
                    chapter=doc.chapter,
                    topic=doc.topic,
                    tags=tags,
                    chunk_count=doc.chunk_count,
                    uploaded_at=doc.uploaded_at,
                    processed_at=doc.processed_at,
                    status=doc.status,
                )
            )

        return filtered_docs

    async def get_document(self, document_id: UUID) -> Optional[Document]:
        """
        Get a document by ID.

        Args:
            document_id: Document ID

        Returns:
            Document if found, None otherwise
        """
        import json
        
        result = await self.db.execute(
            select(DocumentORM).where(DocumentORM.id == str(document_id))
        )
        doc = result.scalar_one_or_none()

        if not doc:
            return None

        # Parse tags from JSON
        tags = []
        if doc.tags:
            try:
                tags = json.loads(doc.tags)
            except json.JSONDecodeError:
                tags = []

        return Document(
            id=UUID(doc.id),
            filename=doc.filename,
            content_type=doc.content_type,
            grade=doc.grade,
            syllabus=doc.syllabus,
            subject=doc.subject,
            chapter=doc.chapter,
            topic=doc.topic,
            tags=tags,
            chunk_count=doc.chunk_count,
            uploaded_at=doc.uploaded_at,
            processed_at=doc.processed_at,
            status=doc.status,
        )

    async def query_by_curriculum(
        self,
        query_text: str,
        filters: Optional[ContentFilters] = None,
        n_results: int = 5,
    ) -> list[dict]:
        """
        Query content by curriculum filters.
        
        Args:
            query_text: The query text to search for
            filters: Optional curriculum filters (grade, syllabus, subject, chapter, topic)
            n_results: Number of results to return
            
        Returns:
            List of matching chunks with metadata and similarity scores
        """
        # Generate query embedding
        query_embedding = self.embedding_service.embed_single(query_text)
        
        # Build where clause from filters
        where = filters.to_chroma_filter() if filters else None
        
        # Query ChromaDB
        results = self.chroma_client.query(
            query_embedding=query_embedding,
            n_results=n_results,
            where=where,
        )
        
        # Format results
        formatted_results = []
        if results and results.get("documents") and results["documents"][0]:
            for i, doc in enumerate(results["documents"][0]):
                metadata = results["metadatas"][0][i] if results.get("metadatas") else {}
                distance = results["distances"][0][i] if results.get("distances") else 0.0
                # Convert distance to similarity (ChromaDB uses cosine distance)
                similarity = 1.0 - distance
                
                formatted_results.append({
                    "content": doc,
                    "metadata": metadata,
                    "similarity": similarity,
                    "grade": metadata.get("grade"),
                    "syllabus": metadata.get("syllabus"),
                    "subject": metadata.get("subject"),
                    "chapter": metadata.get("chapter"),
                    "topic": metadata.get("topic"),
                })
        
        return formatted_results

    async def get_curriculum_summary(self) -> dict:
        """
        Get a summary of available curriculum content.
        
        Returns:
            Dictionary with counts by grade, syllabus, subject, and chapter
        """
        result = await self.db.execute(select(DocumentORM))
        documents = result.scalars().all()
        
        summary = {
            "total_documents": len(documents),
            "by_grade": {},
            "by_syllabus": {},
            "by_subject": {},
            "by_chapter": {},
        }
        
        for doc in documents:
            # Count by grade
            grade_key = str(doc.grade)
            summary["by_grade"][grade_key] = summary["by_grade"].get(grade_key, 0) + 1
            
            # Count by syllabus
            summary["by_syllabus"][doc.syllabus] = summary["by_syllabus"].get(doc.syllabus, 0) + 1
            
            # Count by subject
            summary["by_subject"][doc.subject] = summary["by_subject"].get(doc.subject, 0) + 1
            
            # Count by chapter
            summary["by_chapter"][doc.chapter] = summary["by_chapter"].get(doc.chapter, 0) + 1
        
        return summary
